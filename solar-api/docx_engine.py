#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_engine.py — Word 템플릿 기반 태양광 계약서류 자동입력 엔진
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Word 파일 템플릿의 표(Table) 셀을 직접 찾아 입력값으로 채웁니다.
PDF 오버레이 방식의 좌표 벗어남 문제를 완전히 해소합니다.

지원 서류:
  1. 공급계약신고서
  2. 이용신청서
  3. 전력공급계약신고서
  5. 수금용결제계좌신고서
  6. 전기설비이용계약서
  7. 준법서약서
"""

import copy
import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

# ── 템플릿 디렉터리 ─────────────────────────────────────
TEMPLATE_DIR = Path(__file__).parent / "templates"

# ── 고정 공급사 정보 ────────────────────────────────────
MY = {
    "상호명":           "한화솔루션 주식회사",
    "대표자명":         "박승덕, 남정운, 김동관",
    "대표자명_단독":    "박승덕",
    "주소":             "서울특별시 중구 청계천로 86, 23층 24층(장교동)",
    "사업자번호":       "725-85-01217",
    "법인번호":         "110111-0360935",
    "전화번호":         "02-1600-3400",
    "전기사업자번호":   "서울 2010-가-00001",
    "전기신사업자번호": "2024-서울-0001",
    "업무담당부서":     "한국사업부 채널영업팀",
}


# ════════════════════════════════════════════════════════
#  헬퍼: 셀 텍스트 찾기 / 채우기
# ════════════════════════════════════════════════════════

def _cell_text(cell) -> str:
    """셀의 전체 텍스트 (공백 제거)"""
    return cell.text.strip()


def _set_cell_text(cell, text: str, bold=False, font_size=9, keep_format=True):
    """
    셀의 첫 번째 단락 텍스트를 교체합니다.
    기존 폰트/정렬은 유지하고 텍스트만 바꿉니다.
    """
    if not cell.paragraphs:
        return

    para = cell.paragraphs[0]

    # 기존 런(run)들을 모두 제거하고 새 런 하나만 추가
    # 단, 기존 런의 서식을 복사
    existing_format = None
    if para.runs:
        r = para.runs[0]
        existing_format = {
            'bold':      r.bold,
            'italic':    r.italic,
            'font_name': r.font.name,
            'font_size': r.font.size,
            'color':     r.font.color.rgb if r.font.color and r.font.color.type else None,
        }

    # 기존 런 XML 제거
    for run in para.runs:
        run._element.getparent().remove(run._element)

    # 새 런 추가
    new_run = para.add_run(str(text) if text else "")

    if keep_format and existing_format:
        if existing_format['bold'] is not None:
            new_run.bold = existing_format['bold']
        if existing_format['italic'] is not None:
            new_run.italic = existing_format['italic']
        if existing_format['font_name']:
            new_run.font.name = existing_format['font_name']
        if existing_format['font_size']:
            new_run.font.size = existing_format['font_size']
        if existing_format['color']:
            new_run.font.color.rgb = existing_format['color']
    else:
        new_run.bold = bold
        new_run.font.size = Pt(font_size)

    # 동아시아 폰트 설정 (한글)
    rPr = new_run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')


def _find_cell_after_label(table, label: str, col_offset: int = 1):
    """
    테이블에서 label 텍스트를 가진 셀을 찾고,
    같은 행에서 col_offset 만큼 오른쪽 셀을 반환합니다.
    """
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if label in _cell_text(cell):
                target_ci = ci + col_offset
                if target_ci < len(row.cells):
                    return row.cells[target_ci]
    return None


def _find_cell_by_label(table, label: str):
    """label 텍스트가 포함된 셀 자체를 반환"""
    for row in table.rows:
        for cell in row.cells:
            if label in _cell_text(cell):
                return cell
    return None


def _find_row_by_label(table, label: str):
    """label이 포함된 행(row) 반환"""
    for row in table.rows:
        for cell in row.cells:
            if label in _cell_text(cell):
                return row
    return None


def _get_value_cell(row, label_col: int = 0, value_col: int = None):
    """행에서 값 셀 반환 (value_col이 None이면 마지막 빈 셀 찾기)"""
    cells = row.cells
    if value_col is not None:
        return cells[value_col] if value_col < len(cells) else None
    # 라벨 다음 빈 셀 찾기
    for i in range(label_col + 1, len(cells)):
        if not cells[i].text.strip():
            return cells[i]
    return cells[-1] if cells else None


def _fill_date(doc, year, month, day):
    """
    문서 전체에서 날짜 패턴 '20   년   월   일' 또는 연/월/일 입력란 채우기
    """
    today = datetime.now()
    y = str(year) if year else str(today.year)
    m = str(month) if month else str(today.month)
    d = str(day) if day else str(today.day)

    for para in doc.paragraphs:
        if '년' in para.text and '월' in para.text and '일' in para.text:
            for run in para.runs:
                run.text = run.text.replace(
                    '20   년   월   일',
                    f'20{y[-2:]}년  {m}월  {d}일'
                )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    txt = para.text
                    if '년' in txt and '월' in txt and '일' in txt:
                        for run in para.runs:
                            t = run.text
                            t = re.sub(r'20\s*년\s*월\s*일',
                                       f'20{y[-2:]}년  {m}월  {d}일', t)
                            run.text = t


# ════════════════════════════════════════════════════════
#  서류 1: 공급계약신고서
# ════════════════════════════════════════════════════════

def fill_doc1_supply_contract(d: dict) -> Document:
    """
    1. 태양광발전소 공급계약신고서
    표 구조:
      - 재생에너지전기공급사업자: 고정값 (한화솔루션)
      - 재생에너지발전사업자: 입력값
        · 상호 / 사업자구분
        · 대표자명 / 전화번호(연락처)
        · 주소(사업장주소)
        · 법인등록번호
        · 사업자등록번호
        · 전기사업자등록번호
        · 발전기정보: 발전기주소 / 계량기봉인후입력
        · 설비용량(태양광) / 직접전력거래비율
        · 5년미만철거 / 연간보증공급량
      - 계약일자
      - 서명란
    """
    tpl = TEMPLATE_DIR / "1_공급계약신고서.docx"
    doc = Document(str(tpl))

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            row_text = " ".join(c.text.strip() for c in cells)

            # ── 재생에너지발전사업자 블록 ──
            # 상호
            if any("상" in c.text and "호" in c.text for c in cells[:3]) and \
               any(c.text.strip() == "" for c in cells[3:]):
                # 발전사업자 상호 행인지 확인 (공급사업자 행 제외)
                if "한화솔루션" not in row_text:
                    for i, c in enumerate(cells):
                        t = c.text.strip()
                        if ("상" in t and "호" in t) or t in ("상호", "상 호"):
                            # 다음 값 셀
                            for j in range(i+1, len(cells)):
                                if not cells[j].text.strip():
                                    _set_cell_text(cells[j], d.get("상호명", ""))
                                    break
                            break

            # 대표자명
            if "대표자명" in row_text and "한화솔루션" not in row_text \
               and "박승덕" not in row_text:
                for i, c in enumerate(cells):
                    if "대표자명" in c.text or "대 표 자 명" in c.text:
                        for j in range(i+1, len(cells)):
                            if not cells[j].text.strip():
                                _set_cell_text(cells[j], d.get("대표자명", ""))
                                break
                        # 전화번호 (같은 행 오른쪽)
                        for j in range(i+1, len(cells)):
                            if "전화번호" in cells[j].text or "전 화 번 호" in cells[j].text:
                                for k in range(j+1, len(cells)):
                                    if not cells[k].text.strip():
                                        _set_cell_text(cells[k], d.get("연락처", ""))
                                        break
                                break
                        break

            # 주소 (사업장주소) - 발전사업자
            if "주소" in row_text and "한화솔루션" not in row_text \
               and "서울특별시" not in row_text:
                for i, c in enumerate(cells):
                    if c.text.strip() in ("주소", "주 소") and i < len(cells)-1:
                        for j in range(i+1, len(cells)):
                            if not cells[j].text.strip():
                                addr = d.get("사업장주소", "")
                                _set_cell_text(cells[j], addr)
                                break
                        break

            # 법인등록번호 - 발전사업자
            if "법인" in row_text and "등록번호" in row_text \
               and "110111" not in row_text:
                for i, c in enumerate(cells):
                    if "법인" in c.text and "등록번호" in c.text:
                        for j in range(i+1, len(cells)):
                            if not cells[j].text.strip():
                                _set_cell_text(cells[j], d.get("법인등록번호", ""))
                                break
                        break

            # 사업자등록번호 - 발전사업자
            if "사업자등록번호" in row_text and "725-85" not in row_text:
                for i, c in enumerate(cells):
                    if "사업자등록번호" in c.text:
                        for j in range(i+1, len(cells)):
                            if not cells[j].text.strip():
                                _set_cell_text(cells[j], d.get("사업자등록번호", ""))
                                break
                        break

            # 전기사업자등록번호 - 발전사업자
            if "전기사업자등록번호" in row_text and "서울 2010" not in row_text:
                for i, c in enumerate(cells):
                    if "전기사업자등록번호" in c.text:
                        for j in range(i+1, len(cells)):
                            if not cells[j].text.strip():
                                _set_cell_text(cells[j], d.get("전기사업자등록번호", ""))
                                break
                        break

            # 발전기 주소
            if "발전기" in row_text and ("주소" in row_text or "설치" in row_text):
                for i, c in enumerate(cells):
                    if ("발전기" in c.text and "주소" in c.text) or "발전기주소" in c.text:
                        for j in range(i+1, len(cells)):
                            if not cells[j].text.strip():
                                addr = d.get("발전기주소") or d.get("사업장주소", "")
                                _set_cell_text(cells[j], addr)
                                break
                        break

            # 설비용량 (태양광)
            if "설비용량" in row_text or ("태양광" in row_text and "kW" in row_text):
                for i, c in enumerate(cells):
                    if "태양광" in c.text or "설비용량" in c.text:
                        for j in range(i+1, len(cells)):
                            if not cells[j].text.strip():
                                _set_cell_text(cells[j], str(d.get("설비용량", "")))
                                break
                        break

    # 날짜 채우기
    _fill_date(doc, d.get("계약연도"), d.get("계약월"), d.get("계약일"))

    return doc


# ════════════════════════════════════════════════════════
#  서류 2: 이용신청서
# ════════════════════════════════════════════════════════

def fill_doc2_application(d: dict) -> Document:
    """
    2. 전력거래계약용 이용신청서
    신청인(발전사업자) 블록:
      ①사업자명  ②대표자명
      ③이용장소(발전기주소)  ④고객번호
      ⑤업무담당부서  ⑥전화번호
      ⑦담당자명  ⑧전자우편주소
      ⑨사업허가사항  ⑩희망연계점
    재생에너지전기공급사업자 블록: 고정값
    """
    tpl = TEMPLATE_DIR / "2_이용신청서.docx"
    doc = Document(str(tpl))

    # 필드 매핑: 라벨 → (값, col_offset)
    field_map = {
        "①사 업 자 명":   d.get("상호명", ""),
        "①사업자명":      d.get("상호명", ""),
        "사 업 자 명":    d.get("상호명", ""),
        "②대 표 자 명":  d.get("대표자명", ""),
        "②대표자명":     d.get("대표자명", ""),
        "대 표 자 명":   d.get("대표자명", ""),
        "③이 용 장 소":  d.get("발전기주소") or d.get("사업장주소", ""),
        "③이용장소":     d.get("발전기주소") or d.get("사업장주소", ""),
        "이 용 장 소":   d.get("발전기주소") or d.get("사업장주소", ""),
        "⑥전 화 번 호":  d.get("연락처", ""),
        "⑥전화번호":     d.get("연락처", ""),
        "⑧전자우편주소": d.get("이메일", ""),
    }

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            row_text = " ".join(c.text for c in cells)

            # 재생에너지공급사업자 블록은 건드리지 않음
            if "한화솔루션" in row_text or MY["전화번호"] in row_text:
                continue

            for i, cell in enumerate(cells):
                ct = cell.text.strip()
                for label, value in field_map.items():
                    if label in ct:
                        # 같은 셀에 값이 있는지 확인, 없으면 다음 셀
                        # 셀이 라벨 전용인 경우 바로 옆 셀에 입력
                        if i + 1 < len(cells):
                            target = cells[i + 1]
                            # 값이 아직 비어있으면 채우기
                            if not target.text.strip() or target.text.strip() in ("　", " "):
                                _set_cell_text(target, value)
                        break

    # 날짜
    _fill_date(doc, d.get("계약연도"), d.get("계약월"), d.get("계약일"))
    return doc


# ════════════════════════════════════════════════════════
#  서류 3: 전력공급계약신고서
# ════════════════════════════════════════════════════════

def fill_doc3_power_contract(d: dict) -> Document:
    """
    3. 전력공급계약신고서
    공급사업자 블록: 고정
    발전사업자 블록: 입력
      상호 / 사업자구분
      대표자명 / 전화번호
      주소
      법인등록번호
      사업자등록번호
      전기신사업자등록번호
    발전기 주소 / 설비용량
    계약 기간
    """
    tpl = TEMPLATE_DIR / "3_전력공급계약신고서.docx"
    doc = Document(str(tpl))

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            row_text = " ".join(c.text for c in cells)

            # 공급사업자 블록 스킵
            if "한화솔루션" in row_text or "110111-0360935" in row_text:
                continue

            for i, cell in enumerate(cells):
                ct = cell.text.strip()

                if ct in ("상호", "상 호") and i + 1 < len(cells):
                    if not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("상호명", ""))

                elif ("대표자명" in ct or "대 표 자 명" in ct) and "박승덕" not in row_text:
                    if i + 1 < len(cells) and not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("대표자명", ""))

                elif ("전화번호" in ct or "전 화 번 호" in ct) and "1600" not in row_text:
                    if i + 1 < len(cells) and not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("연락처", ""))

                elif ct in ("주소", "주 소") and "서울특별시" not in row_text:
                    if i + 1 < len(cells) and not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("사업장주소", ""))

                elif "법인" in ct and "등록번호" in ct and "110111" not in row_text:
                    if i + 1 < len(cells) and not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("법인등록번호", ""))

                elif "사업자등록번호" in ct and "725-85" not in row_text:
                    if i + 1 < len(cells) and not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("사업자등록번호", ""))

                elif ("전기신사업자" in ct or "전기사업자등록번호" in ct) \
                        and "서울 2010" not in row_text:
                    if i + 1 < len(cells) and not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("전기사업자등록번호", ""))

    # 발전기주소 / 설비용량 (단락에도 있을 수 있음)
    for para in doc.paragraphs:
        if "발전기" in para.text and "주소" in para.text:
            for run in para.runs:
                if not run.text.strip():
                    run.text = d.get("발전기주소") or d.get("사업장주소", "")

    _fill_date(doc, d.get("계약연도"), d.get("계약월"), d.get("계약일"))
    return doc


# ════════════════════════════════════════════════════════
#  서류 5: 수금용결제계좌신고서
# ════════════════════════════════════════════════════════

def fill_doc5_bank_account(d: dict) -> Document:
    """
    5. 수금용결제계좌신고서
    입력 필드:
      주소 (사업장주소)
      상호
      대표자
      연도/월/일
    """
    tpl = TEMPLATE_DIR / "5_수금용결제계좌신고서.docx"
    doc = Document(str(tpl))

    for para in doc.paragraphs:
        t = para.text
        # "위본인 주 소 :" 줄 처리
        if "위본인" in t or "주 소" in t:
            for run in para.runs:
                if "주 소" in run.text or "주소" in run.text:
                    # 다음 런에 주소값 채우기
                    idx = para.runs.index(run)
                    if idx + 1 < len(para.runs):
                        para.runs[idx+1].text = d.get("사업장주소", "")

        if "상 호" in t or "상호" in t:
            for run in para.runs:
                run.text = run.text.replace(
                    "상	호 :", f"상	호 : {d.get('상호명','')}"
                ).replace(
                    "상 호 :", f"상 호 : {d.get('상호명','')}"
                )

        if "대\t표\t자" in t or "대표자" in t or "대	표	자" in t:
            for run in para.runs:
                if not any(x in run.text for x in ["대", "표", "자", "인"]):
                    run.text = d.get("대표자명", "")

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            row_text = " ".join(c.text for c in cells)

            for i, cell in enumerate(cells):
                ct = cell.text.strip()
                if ("상호" in ct or "상 호" in ct) and i + 1 < len(cells):
                    if not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("상호명", ""))
                elif ("대표자" in ct or "대 표 자" in ct) and i + 1 < len(cells):
                    if not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("대표자명", ""))
                elif ("주소" in ct or "주 소" in ct) and i + 1 < len(cells):
                    if not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("사업장주소", ""))

            # 은행/계좌 정보
            if "은행" in row_text:
                for i, cell in enumerate(cells):
                    if "은행" in cell.text and i + 1 < len(cells):
                        if not cells[i+1].text.strip():
                            _set_cell_text(cells[i+1], d.get("은행명", ""))
            if "예금주" in row_text:
                for i, cell in enumerate(cells):
                    if "예금주" in cell.text and i + 1 < len(cells):
                        if not cells[i+1].text.strip():
                            _set_cell_text(cells[i+1], d.get("예금주명", ""))
            if "계좌번호" in row_text:
                for i, cell in enumerate(cells):
                    if "계좌번호" in cell.text and i + 1 < len(cells):
                        if not cells[i+1].text.strip():
                            _set_cell_text(cells[i+1], d.get("계좌번호", ""))

    _fill_date(doc, d.get("계약연도"), d.get("계약월"), d.get("계약일"))
    return doc


# ════════════════════════════════════════════════════════
#  서류 6: 전기설비이용계약서
# ════════════════════════════════════════════════════════

def fill_doc6_facility_contract(d: dict) -> Document:
    """
    6. 전력거래계약용 전기설비이용계약서
    주요 입력:
      고객(발전사업자) 상호명, 대표자명, 주소, 사업자번호
    """
    tpl = TEMPLATE_DIR / "6_전기설비이용계약서.docx"
    doc = Document(str(tpl))

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            row_text = " ".join(c.text for c in cells)

            # 한국전력거래소 고정값 스킵
            if "한국전력거래소" in row_text or "KPX" in row_text:
                continue
            # 한화솔루션 고정값 스킵
            if "한화솔루션" in row_text:
                continue

            for i, cell in enumerate(cells):
                ct = cell.text.strip()
                if ct in ("상호", "상 호", "사업자명") and i + 1 < len(cells):
                    if not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("상호명", ""))
                elif ("대표자" in ct) and i + 1 < len(cells):
                    if not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("대표자명", ""))
                elif ("주소" in ct) and "서울" not in row_text and i + 1 < len(cells):
                    if not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("사업장주소", ""))
                elif "사업자번호" in ct and "725" not in row_text and i + 1 < len(cells):
                    if not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], d.get("사업자등록번호", ""))

    _fill_date(doc, d.get("계약연도"), d.get("계약월"), d.get("계약일"))
    return doc


# ════════════════════════════════════════════════════════
#  서류 7: 준법서약서
# ════════════════════════════════════════════════════════

def fill_doc7_compliance(d: dict) -> Document:
    """
    7. 준법서약서
    본문 중 [협력업체] → 발전사업자 상호명으로 교체
    서약란: 주소, 상호, 대표자
    """
    tpl = TEMPLATE_DIR / "7_준법서약서.docx"
    doc = Document(str(tpl))

    company = d.get("상호명", "")
    rep = d.get("대표자명", "")
    addr = d.get("사업장주소", "")

    # 본문 단락: [협력업체] 치환
    for para in doc.paragraphs:
        if "[협력업체" in para.text or "협력업체" in para.text:
            for run in para.runs:
                if "[협력업체" in run.text:
                    run.text = run.text.replace("[협력업체", f"[{company}")
                elif "협력업체" == run.text.strip():
                    run.text = company

    # 테이블 내부
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for i, cell in enumerate(cells):
                ct = cell.text.strip()
                # 서약란 주소/상호/대표자
                if ("주소" in ct or "주 소" in ct) and i + 1 < len(cells):
                    if not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], addr)
                elif ("상호" in ct or "상 호" in ct) and i + 1 < len(cells):
                    if not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], company)
                elif ("대표자" in ct or "대 표 자" in ct) and i + 1 < len(cells):
                    if not cells[i+1].text.strip():
                        _set_cell_text(cells[i+1], rep)
                # [협력업체] 치환
                if "[협력업체" in ct:
                    _set_cell_text(cell, ct.replace("[협력업체", f"[{company}"))

    _fill_date(doc, d.get("계약연도"), d.get("계약월"), d.get("계약일"))
    return doc


# ════════════════════════════════════════════════════════
#  통합 생성 함수
# ════════════════════════════════════════════════════════

GENERATORS = {
    "1_공급계약신고서":     fill_doc1_supply_contract,
    "2_이용신청서":         fill_doc2_application,
    "3_전력공급계약신고서": fill_doc3_power_contract,
    "5_수금용결제계좌신고서": fill_doc5_bank_account,
    "6_전기설비이용계약서": fill_doc6_facility_contract,
    "7_준법서약서":         fill_doc7_compliance,
}

TEMPLATE_FILENAMES = {
    "1_공급계약신고서":     "1_공급계약신고서.docx",
    "2_이용신청서":         "2_이용신청서.docx",
    "3_전력공급계약신고서": "3_전력공급계약신고서.docx",
    "5_수금용결제계좌신고서": "5_수금용결제계좌신고서.docx",
    "6_전기설비이용계약서": "6_전기설비이용계약서.docx",
    "7_준법서약서":         "7_준법서약서.docx",
}


def generate_all_docs(d: dict, output_dir: str) -> tuple[list, list]:
    """
    모든 서류를 생성하고 output_dir에 저장합니다.
    반환: (성공 파일 목록, 오류 목록)
    """
    import os
    os.makedirs(output_dir, exist_ok=True)

    # 데이터 기본값 보정
    data = _normalize_data(d)

    success = []
    errors = []

    for name, fn in GENERATORS.items():
        try:
            doc = fn(data)
            safe_company = re.sub(r'[\\/:*?"<>|]', '_', data.get("상호명", "발전소"))
            filename = f"{name}_{safe_company}.docx"
            filepath = os.path.join(output_dir, filename)
            doc.save(filepath)
            success.append(filepath)
        except FileNotFoundError as e:
            errors.append(f"{name}: 템플릿 파일 없음 - {e}")
        except Exception as e:
            errors.append(f"{name}: {e}")

    return success, errors


def _normalize_data(d: dict) -> dict:
    """입력 데이터 정규화 및 기본값 설정"""
    data = dict(d)

    # 기본 날짜
    today = datetime.now()
    if not data.get("계약연도"):
        data["계약연도"] = str(today.year)
    if not data.get("계약월"):
        data["계약월"] = str(today.month)
    if not data.get("계약일"):
        data["계약일"] = str(today.day)

    # 발전기주소 기본값
    if not data.get("발전기주소"):
        data["발전기주소"] = data.get("사업장주소", "")

    # 사업자번호 형식 정규화
    biz = data.get("사업자등록번호", "")
    biz_clean = re.sub(r"[^0-9]", "", biz)
    if len(biz_clean) == 10 and "-" not in biz:
        data["사업자등록번호"] = f"{biz_clean[:3]}-{biz_clean[3:5]}-{biz_clean[5:]}"

    return data


def check_templates() -> dict:
    """템플릿 파일 존재 여부 확인"""
    result = {}
    for name, fname in TEMPLATE_FILENAMES.items():
        path = TEMPLATE_DIR / fname
        result[name] = path.exists()
    return result
